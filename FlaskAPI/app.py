# TO CREATE AND REBUILD VENV ON WINDOWS
# cd FlaskAPI
# virtualenv venv
# .\venv\Scripts\activate
# pip install -r requirements.txt

# TO ADD REQUIREMENTS
# .\venv\Scripts\activate
# pip install package_name
# pip freeze > requirements.txt


import uuid
from text_analysis import FullText, compare_texts
from flask import Flask, request, jsonify, session
import os
import json
import secrets
import openai
import asyncio
from flask_cors import CORS, cross_origin
import uuid
from docx import Document
import firebase_admin
from firebase_admin import credentials, storage, firestore
import io
from datetime import datetime
import nltk
from textblob import TextBlob
from nltk.corpus import stopwords


app = Flask(__name__)

# Initialize Firebase Admin using environment variable or file
firebase_credentials_env = os.environ.get('FIREBASE_CREDENTIALS_JSON')
if firebase_credentials_env:
    cred = credentials.Certificate(json.loads(firebase_credentials_env))
else:
    cred = credentials.Certificate('FlaskAPI/firebasecred.json')

firebase_admin.initialize_app(cred, {'storageBucket': 'traitor-14f52.appspot.com'})
db = firestore.client()
app.config['SESSION_COOKIE_NAME'] = 'session'
app.config['SESSION_TYPE'] = 'filesystem'
app.secret_key = secrets.token_hex(16)
# CORS(app, supports_credentials=True, origins=["DOMAIN"]) PROD MODE | Replace domain with versel or real domain
CORS(app, supports_credentials=True, origins="*")
api_key = os.environ.get('OPENAI_API_KEY', '')
openai.api_key = api_key


@app.before_request
def ensure_session_token():
    print("Before request triggered.")
    if 'sessionToken' not in session:
        print("Token not in session.")
        session['sessionToken'] = str(uuid.uuid4())
        print(f"New sessionToken generated: {session['sessionToken']}")
    else:
        print("Token found in session.")
        print(f"Existing sessionToken found: {session['sessionToken']}")

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

@app.route('/')
def hello_world():
    return 'chatgpt nice brother wojak!'

# This endpoint is called when a user opens the page. It will generate a new session token so they can use the app without logging in.
@app.route('/set-token')
def set_token():
    session['sessionToken'] = str(uuid.uuid4())
    return f"Token set: {session['sessionToken']}"

# This endpoint is returns the created session token.
@app.route('/get-token')
def get_token():
    return f"{session.get('sessionToken', 'No token found.')}"

#This endpoint is the new single route needed to call in order to run ALL tests and then submit the results to the database.
@app.route('/analyze-assignment', methods=['POST'])
async def analyze_assignment():
    # Extract necessary data from request
    data = request.json
    #user_id = data.get('user_id')
    user_id = 'nut'
    session_token = data.get('session_token')
    file_name = data.get('file_name')
    teacherID = data.get('teacherID')
    classID = data.get('classID')
    studentID = data.get('studentID')
    assignmentID = data.get('assignmentID')

    # WILL NEED TO CHANGE TO ACCOUNT BASED SYSTEM AFTER LOGIN IS CREATED
    firebase_file_path = f"files/{session_token}/{file_name}"

    bucket = storage.bucket()
    blob = bucket.blob(firebase_file_path)
    file_blob = blob.download_as_bytes()

    file_stream = io.BytesIO(file_blob)
    doc = Document(file_stream)

    file_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    
    # Start background task for analysis and Firebase update
    print("Starting background task for analysis and Firebase update")
    asyncio.create_task(run_analysis_and_update_db(user_id, file_text, file_stream, teacherID, classID, studentID, assignmentID))
    
    return jsonify({"message": "Analysis in progress. Results will be updated in the database."})

async def run_analysis_and_update_db(user_id, file_text, file_stream, teacherID, classID, studentID, assignmentID):
    
    # Perform analyses (simplified for example)
    print("Performing analyses")
    print("Asking GPT")
    ask_gpt_result = await ask_gpt_analysis(file_text) # Ask GPT if it wrote the text
    print("Reverse Prompt")
    reverse_prompt_result = await reverse_prompt_analysis(file_text) # Reverse engineer the prompt, regenerate the text, and compare the two
    print("File Metadata")
    file_metadata_result = await file_metadata_analysis(file_stream) # Extract metadata from the document and analyze it
    
    # Aggregate results
    analysis_results = {
        "ask_gpt_result": ask_gpt_result,
        "reverse_prompt_result": reverse_prompt_result,
        "file_metadata_result": file_metadata_result,
    }
    
    # Update Firebase with the aggregated results
    #                                     TeacherID                    ClassID                          AssignmentID                     StudentID                    SubmissionID
    # Example Firestore structure: /Users/DFnnwAeWVC4XHqxPOOjf/Classes/JyvUzZ3CrU4OsQPRKsdu/Assignments/ZCN7hs1ZE20EoowIOmOg/Submisssons/xOpjlkFTKgkIV7WvkQqg/Results/RxeDC31VSlTwuzMnDn2K
    print("Sending results to Firebase")
    db.collection('Users').add(analysis_results)
    """resultsRef = db.collection('Users').document(teacherID)\
        .collection('Classes').document(classID)\
        .collection('Assignments').document(assignmentID)\
        .collection('Submissions').document(studentID)\
        .collection('Results')

    resultsDocs = resultsRef.stream()

    for doc in resultsDocs:
        resultsRef.document(doc.id).set(analysis_results)
        break"""

    print("Results sent to Firebase")


async def ask_gpt_analysis(file_text):
    messages = [
        {"role": "system",
         "content": "Please analyze the following text, initially assuming it was written or assisted by a generative AI model like ChatGPT, unless evidence strongly suggests otherwise. Examine the text's structure, style, and content to assess whether it aligns with typical patterns and characteristics of AI-generated text. Consider factors such as language use, complexity, coherence, repetition, and any other relevant aspects that generative AI models tend to exhibit in their writing. After your analysis, provide a concise evaluation, listing the key reasons that support your conclusion on whether the text is likely AI-written or not. Conclude with a definitive answer based on your assessment. After your analysis, please leave a score from 0 to 100 on a new line (with nothing else before or after it indicating its meaning), with a 0 being \"There is no chance that this is AI generated\" and a 100 being \"It is extremely unlikely that this has been written by a human.\". Here is the text for analysis:"},
        {"role": "user", "content": file_text}
    ]

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=messages
        )

        return {
            'testName': 'AskGPT',
            'success': True,
            'response': response['choices'][0]['message']['content'].strip()
        }

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }
    
async def reverse_prompt_analysis(file_text):
    # Reverse engineer the prompt
    description_messages = [
        {"role": "user",
         "content": "Given the following text input, please analyze and infer the most likely prompt that could have been used to generate this text with a ChatGPT-like model. Assume that the text was either written by or assisted by a model similar to ChatGPT. Based on your analysis, reconstruct the original user prompt as accurately as possible. Additionally, your response should match the length and style of the provided text to ensure a coherent and contextually appropriate output. Do not reply anything thing else other than the prompt response, you do not need to include anymore information other than the prompt. You should include a paragraph or character count in your prompt to have a text of the same length. Here is the text input:\n\n" + file_text}
    ]

    try:
        description_response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=description_messages
        )
        description = description_response['choices'][0]['message']['content'].strip()

        reverse_messages = [
            {"role": "user", "content": description}
        ]

        reverse_response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=reverse_messages
        )
        reversed_prompt = reverse_response['choices'][0]['message']['content'].strip()

        # Compare the original text with the regenerated text
        analysis1 = FullText(file_text)
        analysis2 = FullText(reversed_prompt)
        comparison_results = compare_texts(analysis1, analysis2)

        return {
            'testName': 'ReversePrompt',
            'success': True,
            'original_prompt': file_text,
            'reversed_prompt': reversed_prompt,
            'description': description,
            'comparison_results': comparison_results
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }
    
async def file_metadata_analysis(file_stream):
    doc = Document(file_stream)

    core_properties = doc.core_properties
    attributes = [
        'title', 'author', 'created', 'modified', 'last_modified_by',
        'description', 'category', 'comments', 'subject', 'keywords',
        'version', 'revision', 'identifier', 'language', 'content_status'
    ]

    metadata = {}
    for attr in attributes:
        if hasattr(core_properties, attr):
            value = getattr(core_properties, attr)
            if value:
                metadata[attr] = str(value)

    metadata_prompt = json.dumps(metadata, indent=2)
    
    messages = [
        {"role": "system", 
         "content": "Please analyze the following Word document metadata for any odd or suspicious characteristics that may indicate cheating. Examine the metadata's structure, content, and any other relevant aspects. After your analysis, provide a concise evaluation, listing the key reasons that support your conclusion. After your analysis, please leave a score from 0 to 100 on a new line (with nothing else before or after it indicating its meaning), with a 0 being \"There is no chance that this is AI generated\" and a 100 being \"It is extremely unlikely that this has been written by a human.\". Include a formatted copy of the metadata at the top of your response."},
        {"role": "user", "content": metadata_prompt}
    ]

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=messages
        )

        return {
            'testName': 'MetadataAnalysis',
            'success': True,
            'response': response['choices'][0]['message']['content'].strip()
        }

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }
############################################################################################################################################################################
# This endpoint is one of our detection methods. It will ask GPT if it wrote the text or not and try to provide reasoning.
@app.route('/askgpt', methods=['POST'])
def askGPT():
    data = request.json
    user_prompt = data.get('prompt', '')

    messages = [
        {"role": "system",
         "content": "Please analyze the following text, initially assuming it was written or assisted by a generative AI model like ChatGPT, unless evidence strongly suggests otherwise. Examine the text's structure, style, and content to assess whether it aligns with typical patterns and characteristics of AI-generated text. Consider factors such as language use, complexity, coherence, repetition, and any other relevant aspects that generative AI models tend to exhibit in their writing. After your analysis, provide a concise evaluation, listing the key reasons that support your conclusion on whether the text is likely AI-written or not. Conclude with a definitive answer based on your assessment. After your analysis, please leave a score from 0 to 100 on a new line (with nothing else before or after it indicating its meaning), with a 0 being \"There is no chance that this is AI generated\" and a 100 being \"It is extremely unlikely that this has been written by a human.\". Here is the text for analysis:"},
        {"role": "user", "content": user_prompt}
    ]

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=messages
        )

        return jsonify({
            'testName': 'AskGPT',
            'success': True,
            'response': response['choices'][0]['message']['content'].strip()
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

# This endpoint is another detection method. It will ask GPT to reverse engineer the prompt that was used to generate the text.
# It will then ask GPT to regenerate the text using the reverse engineered prompt.
@app.route('/reverseprompt', methods=['POST'])
def reversePrompt():
    data = request.json
    user_prompt = data.get('prompt', '')

    description_messages = [
        {"role": "user",
         "content": "Given the following text input, please analyze and infer the most likely prompt that could have been used to generate this text with a ChatGPT-like model. Assume that the text was either written by or assisted by a model similar to ChatGPT. Based on your analysis, reconstruct the original user prompt as accurately as possible. Additionally, your response should match the length and style of the provided text to ensure a coherent and contextually appropriate output. Do not reply anything thing else other than the prompt response, you do not need to include anymore information other than the prompt. You should include a paragraph or character count in your prompt to have a text of the same length. Here is the text input:\n\n" + user_prompt}
    ]

    try:
        description_response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=description_messages
        )
        description = description_response['choices'][0]['message']['content'].strip()

        reverse_messages = [
            {"role": "user", "content": description}
        ]

        reverse_response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=reverse_messages
        )
        reversed_prompt = reverse_response['choices'][0]['message']['content'].strip()

        return jsonify({
            'testName': 'ReversePrompt',
            'success': True,
            'original_prompt': user_prompt,
            'reversed_prompt': reversed_prompt,
            'description': description
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

# This endpoint is called when the user drags in a file. It will download the file from the Firebase Filebucket, extract the file's metadata
# and then send the metadata to GPT to analyze it and provide the analysis.
@app.route('/documentscan', methods=['POST'])
def document_scan():
    data = request.json
    session_token = data.get('session_token')
    file_name = data.get('file_name')

    firebase_file_path = f"files/{session_token}/{file_name}"

    bucket = storage.bucket()
    blob = bucket.blob(firebase_file_path)
    file_blob = blob.download_as_bytes()

    file_stream = io.BytesIO(file_blob)
    metadata, full_text, analysis_result = extract_metadata_and_text(file_stream)

    return jsonify({"metadata": metadata, "text": full_text, "analysis": analysis_result})

def extract_metadata_and_text(file_stream):
    doc = Document(file_stream)

    core_properties = doc.core_properties
    attributes = [
        'title', 'author', 'created', 'modified', 'last_modified_by',
        'description', 'category', 'comments', 'subject', 'keywords',
        'version', 'revision', 'identifier', 'language', 'content_status'
    ]

    metadata = {}
    for attr in attributes:
        if hasattr(core_properties, attr):
            value = getattr(core_properties, attr)
            if value:
                metadata[attr] = str(value)

    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    analysis_result = analyze_metadata_with_chatgpt(metadata)

    return metadata, full_text, analysis_result

def analyze_metadata_with_chatgpt(metadata):
    metadata_prompt = json.dumps(metadata, indent=2)
    
    messages = [
        {"role": "system", 
         "content": "Please analyze the following Word document metadata for any odd or suspicious characteristics that may indicate cheating. Examine the metadata's structure, content, and any other relevant aspects. After your analysis, provide a concise evaluation, listing the key reasons that support your conclusion. After your analysis, please leave a score from 0 to 100 on a new line (with nothing else before or after it indicating its meaning), with a 0 being \"There is no chance that this is AI generated\" and a 100 being \"It is extremely unlikely that this has been written by a human.\". Include a formatted copy of the metadata at the top of your response."},
        {"role": "user", "content": metadata_prompt}
    ]

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=messages
        )

        return {
            'testName': 'MetadataAnalysis',
            'success': True,
            'response': response['choices'][0]['message']['content'].strip()
        }

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

# This endpoint is called when the user drags in a file. 
# It will upload the file to Firebase Storage and provide the text to display the text in the textbox.
@app.route('/extract-text', methods=['POST'])
def extract_text():
    try:
        print("extract_text - Started processing request")
        data = request.json
        session_token = data.get('session_token')
        file_name = data.get('file_name')

        print(f"extract_text - Session Token: {session_token}, File Name: {file_name}")

        if not session_token or not file_name:
            print("extract_text - Missing session token or file name")
            return jsonify({"error": "Missing session token or file name"}), 400

        # WILL NEED TO CHANGE TO ACCOUNT BASED SYSTEM AFTER LOGIN IS CREATED
        firebase_file_path = f"files/{session_token}/{file_name}"
        full_firebase_file_path = f"gs://traitor-14f52.appspot.com/{firebase_file_path}"

        # Push the file path to the Firestore with student's assignment
        #                                     TeacherID                    ClassID                          AssignmentID                     StudentID                    SubmissionID
        # Example Firestore structure: /Users/DFnnwAeWVC4XHqxPOOjf/Classes/JyvUzZ3CrU4OsQPRKsdu/Assignments/ZCN7hs1ZE20EoowIOmOg/Submisssons/xOpjlkFTKgkIV7WvkQqg/Results/RxeDC31VSlTwuzMnDn2K
        db.collection('Users').document('DFnnwAeWVC4XHqxPOOjf')\
            .collection('Classes').document('JyvUzZ3CrU4OsQPRKsdu')\
            .collection('Assignments').document('cZzd0pouqJCAmVxqdlk4')\
            .collection('Submissions').add({'assignmentTitle': file_name, 'bucketFilePath': full_firebase_file_path, 'submissionDate': datetime.now()})
        print(f"extract_text - Firebase File Path: {firebase_file_path}")

        bucket = storage.bucket()
        blob = bucket.blob(firebase_file_path)
        file_blob = blob.download_as_bytes()

        print("extract_text - File blob downloaded")

        file_stream = io.BytesIO(file_blob)
        doc = Document(file_stream)

        full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

        print("extract_text - Text extracted successfully")

        return jsonify({"text": full_text})
    except Exception as e:
        print(f"extract_text - Error: {str(e)}")
        return jsonify({"error": str(e)}), 500

# This endpoint is another detection method. It will take the original text and reverese GPT'd text and
# then compare the two texts to see if they are similar or not through the text_analysis python file.
@app.route('/analyze-compare-texts', methods=['POST'])
def analyze_compare_texts():
    data = request.json
    userText = data.get('text1')
    generatedText = data.get('text2')

    if not userText or not generatedText:
        return jsonify({"error": "Both text1 and text2 are required"}), 400

    try:
        analysis1 = FullText(userText)
        analysis2 = FullText(generatedText)
        comparison_results = compare_texts(analysis1, analysis2)

        return jsonify(comparison_results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# This endpoint was just to test the Firestore, not really necessary now but keeping it just in case
@app.route('/test-db', methods=['GET'])
def testDB():
    try:
        doc_ref = db.collection('test').document('test')
        doc_ref.set({
            'test': 'test'
        })

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True)
